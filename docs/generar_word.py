#!/usr/bin/env python3
"""
Genera documentacion profesional Word (.docx) para EcoMarket - UNAP Puno
"""
import os
import zipfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colores ───────────────────────────────────────────────────
C_GREEN      = RGBColor(27,  94,  32)
C_DKGREEN    = RGBColor(10,  60,  15)
C_DARK       = RGBColor(30,  30,  30)
C_GRAY       = RGBColor(90,  90,  90)
C_WHITE      = RGBColor(255, 255, 255)
HEX_GREEN    = "1B5E20"
HEX_LTGREEN  = "E8F5E9"
HEX_WHITE    = "FFFFFF"

# ── Utilidades ────────────────────────────────────────────────
def font(run, size=11, bold=False, italic=False, color=C_DARK, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color

def para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        r = p.add_run(text)
        font(r)
    return p

def hline(doc, color=HEX_GREEN, width="12"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ["top", "bottom"]:
        bd = OxmlElement(f"w:{side}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), width)
        bd.set(qn("w:space"), "1")
        bd.set(qn("w:color"), color)
        pBdr.append(bd)
    pPr.append(pBdr)

def shading(cell, fill_hex):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tc.append(shd)

def heading(doc, text, level=1, sb=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    if level == 1:
        font(r, size=14, bold=True, color=C_GREEN)
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bd = OxmlElement("w:bottom")
        bd.set(qn("w:val"), "single"); bd.set(qn("w:sz"), "6")
        bd.set(qn("w:space"), "1");   bd.set(qn("w:color"), HEX_GREEN)
        pBdr.append(bd); pPr.append(pBdr)
    elif level == 2:
        font(r, size=12, bold=True, color=C_DARK)
    else:
        font(r, size=11, bold=True, color=C_GRAY)

def qa(doc, question, answer):
    p_q = doc.add_paragraph()
    p_q.paragraph_format.left_indent = Cm(0.6)
    p_q.paragraph_format.space_before = Pt(6)
    p_q.paragraph_format.space_after  = Pt(2)
    r = p_q.add_run(f"P: {question}")
    font(r, bold=True, color=C_GREEN, size=10)

    p_a = doc.add_paragraph()
    p_a.paragraph_format.left_indent = Cm(0.6)
    p_a.paragraph_format.space_before = Pt(2)
    p_a.paragraph_format.space_after  = Pt(8)
    r = p_a.add_run(f"R: {answer}")
    font(r, color=C_DARK, size=10)

def use_case(doc, code, name, actors, desc, pre, flow, post):
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    tbl = doc.add_table(rows=6, cols=2)
    tbl.style = "Table Grid"

    # Header row (merged)
    hrow = tbl.rows[0]
    hrow.cells[0].merge(hrow.cells[1])
    hrow.cells[0].paragraphs[0].clear()
    r = hrow.cells[0].paragraphs[0].add_run(f"{code}  ·  {name}")
    font(r, bold=True, color=C_WHITE, size=10)
    hrow.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading(hrow.cells[0], HEX_GREEN)

    # Data rows
    labels = ["Actores", "Descripción", "Precondición", "Flujo principal", "Postcondición"]
    values = [actors, desc, pre, flow, post]
    for i, (lbl, val) in enumerate(zip(labels, values)):
        row = tbl.rows[i + 1]
        lc, vc = row.cells[0], row.cells[1]
        lc.paragraphs[0].clear()
        r = lc.paragraphs[0].add_run(lbl)
        font(r, bold=True, size=9, color=C_DKGREEN)
        shading(lc, HEX_LTGREEN)
        vc.paragraphs[0].clear()
        r = vc.paragraphs[0].add_run(val)
        font(r, size=9, color=C_DARK)

    # Column widths
    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            cell.width = Cm(4) if j == 0 else Cm(12)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def info_table(doc, rows_data):
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = "Table Grid"
    for i, (lbl, val) in enumerate(rows_data):
        row = tbl.rows[i]
        row.cells[0].paragraphs[0].clear()
        r = row.cells[0].paragraphs[0].add_run(lbl)
        font(r, bold=True, size=9, color=C_GREEN)
        shading(row.cells[0], HEX_LTGREEN)
        row.cells[1].paragraphs[0].clear()
        r = row.cells[1].paragraphs[0].add_run(val)
        font(r, size=9, color=C_DARK)
        for cell in row.cells:
            cell.width = Cm(4.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_footer(doc):
    for section in doc.sections:
        fp = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.clear()
        r1 = fp.add_run("EcoMarket — UNAP Puno   |   Página ")
        font(r1, size=8, italic=True, color=C_GRAY)
        for tag, txt in [("begin", ""), ("", " PAGE "), ("end", "")]:
            if tag in ("begin", "end"):
                fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), tag)
                rx = fp.add_run(); rx._r.append(fc)
            else:
                it = OxmlElement("w:instrText"); it.text = txt
                fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
                fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "end")
                rx = fp.add_run(); rx._r.append(fc1); rx._r.append(it); rx._r.append(fc2)
        font(rx, size=8, color=C_GRAY)

# ════════════════════════════════════════════════════════════
# DOCUMENTO
# ════════════════════════════════════════════════════════════
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

add_footer(doc)

# ── CARÁTULA ─────────────────────────────────────────────────
for _ in range(7):
    doc.add_paragraph()

p = para(doc, "UNIVERSIDAD NACIONAL DEL ALTIPLANO PUNO", WD_ALIGN_PARAGRAPH.CENTER, sa=4)
font(p.runs[0], size=15, bold=True, color=C_DARK)
p = para(doc, "Facultad de Ingeniería de Sistemas", WD_ALIGN_PARAGRAPH.CENTER, sa=2)
font(p.runs[0], size=12, color=C_DARK)
p = para(doc, "Escuela Profesional de Ingeniería de Sistemas", WD_ALIGN_PARAGRAPH.CENTER, sa=10)
font(p.runs[0], size=11, color=C_GRAY)

hline(doc, HEX_GREEN, "18")

for _ in range(2):
    doc.add_paragraph()

p = para(doc, "EcoMarket", WD_ALIGN_PARAGRAPH.CENTER, sa=6)
font(p.runs[0], size=40, bold=True, color=C_GREEN)

p = para(doc, "Plataforma de Comercio Electrónico de Productos Ecológicos", WD_ALIGN_PARAGRAPH.CENTER, sa=3)
font(p.runs[0], size=14, color=C_DARK)
p = para(doc, "con Auditoría Química mediante Inteligencia Artificial", WD_ALIGN_PARAGRAPH.CENTER, sa=14)
font(p.runs[0], size=14, color=C_DARK)

p = para(doc, "Documento de Análisis de Requerimientos", WD_ALIGN_PARAGRAPH.CENTER, sa=2)
font(p.runs[0], size=12, italic=True, color=C_GRAY)
p = para(doc, "Entrevista a Actores Clave · Especificación de Casos de Uso", WD_ALIGN_PARAGRAPH.CENTER, sa=16)
font(p.runs[0], size=12, italic=True, color=C_GRAY)

hline(doc, HEX_GREEN, "18")
for _ in range(3):
    doc.add_paragraph()

date_str = datetime.now().strftime("%-d de %B de %Y") if os.name != "nt" else datetime.now().strftime("%d de %B de %Y")

cover_info = [
    ("Curso:",       "Ingeniería de Software"),
    ("Grupo:",       "Grupo A — Octavo Semestre"),
    ("Docente:",     "Ing. Consorcio de Ingeniería de Software"),
    ("Integrantes:", "Josep Vladimir · Miembros del equipo EcoMarket"),
    ("Fecha:",       date_str),
]
info_table(doc, cover_info)

doc.add_page_break()

# ── ÍNDICE ───────────────────────────────────────────────────
heading(doc, "Tabla de Contenidos", 1, sb=6)

toc = [
    ("1.",    "Introducción"),
    ("1.1.",  "Contexto del Proyecto"),
    ("1.2.",  "Objetivos del Sistema"),
    ("2.",    "Entrevista a Actores Clave"),
    ("2.1.",  "Metodología de Entrevistas"),
    ("2.2.",  "Entrevista a Productor Ecológico — Juan Quispe"),
    ("2.3.",  "Entrevista a Consumidor Local — María Mamani"),
    ("2.4.",  "Entrevista a Comerciante — Rosa Condori"),
    ("2.5.",  "Conclusiones de la Entrevista"),
    ("3.",    "Casos de Uso del Sistema"),
    ("3.1.",  "Actores del Sistema"),
    ("3.2.",  "CU-01 · Registro de Usuario"),
    ("3.3.",  "CU-02 · Inicio de Sesión"),
    ("3.4.",  "CU-03 · Gestión de Catálogo de Productos"),
    ("3.5.",  "CU-04 · Auditoría Química de Producto"),
    ("3.6.",  "CU-05 · Compra de Productos"),
    ("3.7.",  "CU-06 · Ver Catálogo y Filtrar"),
    ("3.8.",  "CU-07 · Gestión de Usuarios"),
    ("3.9.",  "CU-08 · Generación de Certificado PDF"),
    ("3.10.", "CU-09 · Dashboard de Ventas"),
    ("3.11.", "CU-10 · Notificaciones Push"),
    ("4.",    "Glosario de Términos"),
]
for num, item in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    is_main = num.count(".") == 1 and num[-1] == "."
    r1 = p.add_run(f"{num:<7}")
    font(r1, bold=is_main, size=10, color=C_GREEN if is_main else C_GRAY)
    r2 = p.add_run(item)
    font(r2, bold=is_main, size=10, color=C_DARK if is_main else C_DARK)

doc.add_page_break()

# ── 1. INTRODUCCIÓN ──────────────────────────────────────────
heading(doc, "1. Introducción", 1)
heading(doc, "1.1 Contexto del Proyecto", 2)

body1 = (
    "EcoMarket es una plataforma de comercio electrónico diseñada para la región de Puno que conecta "
    "a productores locales de productos ecológicos y orgánicos con consumidores conscientes del medio ambiente. "
    "La región de Puno, ubicada en el altiplano peruano a orillas del lago Titicaca, posee una rica tradición "
    "agrícola basada en cultivos nativos como la quinua, cañihua, kiwicha y papas nativas, producidos de manera "
    "tradicional y ecológica por pequeños agricultores que enfrentan serias dificultades para acceder a "
    "mercados más amplios."
)
p = para(doc, sa=8); r = p.add_run(body1); font(r, color=C_DARK)

body2 = (
    "EcoMarket surge como una solución integral que no solo facilita la venta en línea, sino que incorpora "
    "un sistema de auditoría química basado en Inteligencia Artificial que analiza los ingredientes de cada "
    "producto y emite certificados de pureza ecológica con puntuaciones Eco-Score, generando confianza tanto "
    "en compradores locales como en mercados externos e internacionales."
)
p = para(doc, sa=8); r = p.add_run(body2); font(r, color=C_DARK)

heading(doc, "1.2 Objetivos del Sistema", 2)

p = para(doc, sa=3); r = p.add_run("Objetivo General"); font(r, bold=True, color=C_GREEN)
p = para(doc, sa=10)
r = p.add_run(
    "Desarrollar una plataforma de comercio electrónico para productos ecológicos de la región de Puno "
    "que integre un sistema de auditoría química mediante IA para garantizar la autenticidad y calidad "
    "de los productos ofertados."
)
font(r, color=C_DARK)

p = para(doc, sa=3); r = p.add_run("Objetivos Específicos"); font(r, bold=True, color=C_GREEN)

specific = [
    "Diseñar un catálogo digital con filtros por categoría, Eco-Score y ubicación geográfica.",
    "Desarrollar un módulo de autenticación seguro con roles: Consumidor, Productor y Administrador.",
    "Implementar un motor de auditoría química con IA que calcule el Eco-Score de cada producto.",
    "Integrar pagos locales (Yape, Plin) e internacionales (Stripe).",
    "Generar certificados PDF de auditoría con hash SHA-256 para garantizar la inmutabilidad.",
    "Diseñar una arquitectura de microservicios escalable y desplegable en la nube.",
]
for obj in specific:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(obj); font(r, size=10, color=C_DARK)

doc.add_page_break()

# ── 2. ENTREVISTA ────────────────────────────────────────────
heading(doc, "2. Entrevista a Actores Clave", 1)
heading(doc, "2.1 Metodología de Entrevistas", 2)

metodo = (
    "Se realizaron entrevistas semiestructuradas a tres actores clave del ecosistema de productos ecológicos "
    "en la región de Puno: un productor agrícola, un consumidor local y una comerciante del mercado central "
    "de la ciudad de Puno. Las entrevistas se llevaron a cabo en marzo de 2026 con el objetivo de identificar "
    "las necesidades, expectativas y requerimientos funcionales de la plataforma EcoMarket."
)
p = para(doc, sa=10); r = p.add_run(metodo); font(r, color=C_DARK)

# ── Entrevista 1 ─────────────────────────────────────────────
heading(doc, "2.2 Entrevista al Productor Ecológico", 2)
info_table(doc, [
    ("Entrevistado:", "Juan Quispe Callata"),
    ("Edad:",         "52 años"),
    ("Ocupación:",    "Agricultor ecológico — Comunidad de Chucuito, Puno"),
    ("Productos:",    "Quinua real, cañihua, papas nativas"),
    ("Fecha:",        "5 de marzo de 2026"),
])

qa(doc,
   "Don Juan, ¿cómo vende actualmente sus productos ecológicos?",
   "Vendo en la feria sabatina de Puno y a veces a intermediarios que llevan mis productos a Juliaca y "
   "Arequipa. El problema es que los intermediarios me pagan muy poco. La quinua real me compran a 3 soles "
   "el kilo y ellos la venden a 12 soles en la ciudad. No tengo forma de llegar directamente al consumidor final.")

qa(doc,
   "¿Qué opina de vender sus productos por internet?",
   "Me gustaría, pero no sé cómo hacerlo. Mis hijos me enseñaron a usar WhatsApp, pero una tienda virtual "
   "me parece complicado. Además, la gente desconfía, quieren saber si realmente es orgánico. Yo tengo mi "
   "certificación del SENASA, pero no sé cómo mostrarla en internet.")

qa(doc,
   "¿Qué tan importante es que sus productos tengan certificación ecológica?",
   "Es muy importante. Yo no uso pesticidas ni fertilizantes químicos, todo es natural como se ha hecho "
   "siempre en mi comunidad. Si la plataforma pudiera analizar mis productos y dar un puntaje de confianza, "
   "sería bueno para diferenciarme de los que no son genuinos.")

qa(doc,
   "¿Qué funcionalidades necesita en la plataforma?",
   "Primero, que sea fácil de usar desde mi celular. Segundo, que los clientes vean mi certificación y el "
   "análisis de mis productos. Tercero, recibir pagos directos sin intermediarios. Y cuarto, que me ayude "
   "a coordinar las entregas, porque a veces vendo a gente de fuera de Puno.")

# ── Entrevista 2 ─────────────────────────────────────────────
heading(doc, "2.3 Entrevista a la Consumidora Local", 2)
info_table(doc, [
    ("Entrevistado:", "María Mamani Huanca"),
    ("Edad:",         "28 años"),
    ("Ocupación:",    "Profesora de nivel primario"),
    ("Lugar:",        "Puno ciudad"),
    ("Fecha:",        "8 de marzo de 2026"),
])

qa(doc,
   "María, ¿dónde compra actualmente sus alimentos?",
   "Compro en el mercado central y en algunas bodegas del barrio. Me gustaría comprar productos ecológicos "
   "porque he leído que son más saludables, pero no siempre sé cuál es realmente orgánico. Todo se ve igual "
   "y a veces los precios son más altos sin saber si vale la pena.")

qa(doc,
   "¿Le gustaría una plataforma con productos ecológicos certificados?",
   "Sí, definitivamente. Sobre todo si puedo ver el análisis de cada producto: su puntuación ecológica y "
   "qué ingredientes tiene. También me gustaría saber de dónde viene, porque prefiero apoyar a los "
   "agricultores de la región. Con entregas a domicilio sería perfecto.")

qa(doc,
   "¿Qué métodos de pago prefiere?",
   "Uso Yape y Plin casi a diario. También tengo tarjeta de débito. Sería bueno tener varias opciones. "
   "A veces me da desconfianza poner datos de tarjeta en internet, pero si la plataforma tiene buena "
   "reputación lo usaría sin problemas.")

# ── Entrevista 3 ─────────────────────────────────────────────
heading(doc, "2.4 Entrevista a la Comerciante del Mercado", 2)
info_table(doc, [
    ("Entrevistado:", "Rosa Condori Vilca"),
    ("Edad:",         "45 años"),
    ("Ocupación:",    "Comerciante — Mercado Central de Puno"),
    ("Productos:",    "Quinuas, habas, cebada, quesos regionales"),
    ("Fecha:",        "10 de marzo de 2026"),
])

qa(doc,
   "Sra. Rosa, ¿ha considerado vender sus productos en línea?",
   "Lo he pensado, pero no tengo tiempo para aprender plataformas complicadas. El envío también es un "
   "problema porque vendo productos frescos. Me preocupa que los clientes no paguen o me estafen. "
   "Prefiero el trato directo, cara a cara.")

qa(doc,
   "¿Cree que una plataforma digital aumentaría sus ventas?",
   "Si es fácil de usar y tiene buena reputación, sí. Sobre todo si atrae a turistas o gente de otras "
   "regiones que busca productos típicos de calidad. Muchos turistas preguntan por quinua real, cañihua "
   "y quesos de cabra, pero no siempre encuentran dónde comprar con confianza.")

qa(doc,
   "¿Qué opina de un sistema de auditoría de ingredientes en la plataforma?",
   "Me parece excelente. Así el cliente puede ver que el producto es realmente orgánico y no solo "
   "publicidad. Yo conozco personalmente a los productores de los que compro. Si la plataforma puede "
   "certificar eso con un análisis, sería un gran respaldo para los comerciantes honestos.")

# ── Conclusiones ─────────────────────────────────────────────
heading(doc, "2.5 Conclusiones de la Entrevista", 2)

conclusiones = [
    ("Necesidad de digitalización:",
     "Los productores dependen de intermediarios que reducen sus márgenes. "
     "Una plataforma directa elimina esta dependencia y aumenta la rentabilidad del productor."),
    ("Confianza y certificación:",
     "Los consumidores demandan mecanismos verificables de autenticidad. "
     "El Eco-Score y los certificados PDF con SHA-256 son el diferenciador clave del sistema."),
    ("Facilidad de uso:",
     "La interfaz debe ser intuitiva y accesible desde celulares de gama media, "
     "considerando que muchos productores tienen familiaridad básica con aplicaciones móviles."),
    ("Pagos locales indispensables:",
     "Yape y Plin son los métodos de pago preferidos en Puno. "
     "Stripe es necesario para compradores internacionales y turistas."),
    ("Valoración del origen local:",
     "Mostrar la procedencia geográfica y las prácticas tradicionales de cultivo "
     "genera confianza y diferencia los productos de EcoMarket de los de otras plataformas."),
    ("Logística de entrega:",
     "Se requiere un sistema de coordinación que contemple ventas locales (Puno) "
     "y envíos a nivel nacional con tracking básico."),
]
for term, desc in conclusiones:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"{term} "); font(r1, bold=True, color=C_GREEN, size=10)
    r2 = p.add_run(desc);       font(r2, color=C_DARK, size=10)

doc.add_page_break()

# ── 3. CASOS DE USO ──────────────────────────────────────────
heading(doc, "3. Casos de Uso del Sistema", 1)
heading(doc, "3.1 Actores del Sistema", 2)

# Actors table
atbl = doc.add_table(rows=5, cols=2)
atbl.style = "Table Grid"
for cell in atbl.rows[0].cells:
    shading(cell, HEX_GREEN)
atbl.rows[0].cells[0].paragraphs[0].clear()
atbl.rows[0].cells[1].paragraphs[0].clear()
r = atbl.rows[0].cells[0].paragraphs[0].add_run("Actor"); font(r, bold=True, color=C_WHITE, size=10)
r = atbl.rows[0].cells[1].paragraphs[0].add_run("Descripción"); font(r, bold=True, color=C_WHITE, size=10)

actors_data = [
    ("Cliente",         "Usuario final: navega el catálogo, compra productos y visualiza certificados de auditoría."),
    ("Productor",       "Vendedor que gestiona su catálogo ecológico, visualiza auditorías y estadísticas de venta."),
    ("Administrador",   "Gestiona usuarios, ejecuta auditorías químicas, aprueba/rechaza productos y genera reportes."),
    ("Sistema",         "Actor secundario que ejecuta notificaciones push y auditorías automáticas."),
]
for i, (act, dsc) in enumerate(actors_data):
    row = atbl.rows[i + 1]
    row.cells[0].paragraphs[0].clear()
    r = row.cells[0].paragraphs[0].add_run(act); font(r, bold=True, size=9, color=C_GREEN)
    shading(row.cells[0], HEX_LTGREEN)
    row.cells[1].paragraphs[0].clear()
    r = row.cells[1].paragraphs[0].add_run(dsc); font(r, size=9, color=C_DARK)
    row.cells[0].width = Cm(3.5); row.cells[1].width = Cm(12.5)

doc.add_paragraph().paragraph_format.space_after = Pt(8)
heading(doc, "3.2 Especificación de Casos de Uso", 2)

# ── Use cases ────────────────────────────────────────────────
use_case(doc,
    "CU-01", "Registro de Usuario",
    "Cliente, Productor",
    "Permite a un nuevo usuario registrarse. Los Clientes ingresan nombre, email, contraseña, teléfono y dirección. "
    "Los Productores añaden RUC, representante legal y certificaciones ecológicas.",
    "El email no debe estar registrado previamente en el sistema.",
    "1. Accede al formulario de registro. 2. Selecciona rol (Cliente o Productor). "
    "3. Completa campos requeridos. 4. Acepta términos y condiciones. "
    "5. Sistema valida datos y crea la cuenta. 6. Envía correo de bienvenida y retorna token JWT.",
    "Usuario registrado y autenticado. Registro en BD con estado 'activo' y token JWT válido por 24 h.")

use_case(doc,
    "CU-02", "Inicio de Sesión",
    "Cliente, Productor, Administrador",
    "Permite a un usuario registrado autenticarse mediante email y contraseña, obteniendo un token JWT "
    "con los permisos de su rol para acceder a recursos protegidos.",
    "El usuario debe estar registrado y activo en el sistema.",
    "1. Ingresa email y contraseña. 2. Sistema valida credenciales contra BD. "
    "3. Si son correctas genera token JWT con roles y permisos. 4. Retorna token y datos del usuario.",
    "Token JWT válido por 24 horas que autoriza acceso a recursos según el rol del usuario.")

use_case(doc,
    "CU-03", "Gestión de Catálogo de Productos",
    "Productor",
    "Permite al productor crear, editar y gestionar sus productos: imágenes, ingredientes, "
    "datos nutricionales, origen GPS, fechas de producción/vencimiento y precio.",
    "Productor autenticado con RUC verificado por el Administrador.",
    "1. Accede al panel de gestión. 2. Selecciona 'Nuevo Producto' o edita uno existente. "
    "3. Completa: nombre, descripción, categoría, precio, stock, imágenes, ingredientes, coordenadas GPS. "
    "4. Sistema registra con estado 'PENDIENTE'. 5. Producto visible solo para el productor hasta ser auditado.",
    "Producto en BD con estado 'PENDIENTE'. Administrador notificado para revisión.")

use_case(doc,
    "CU-04", "Auditoría Química de Producto",
    "Administrador",
    "El Administrador ejecuta el motor de IA sobre un producto pendiente. El sistema analiza cada ingrediente "
    "contra la base de datos química, calcula Eco-Score (0–100), asigna badges y genera certificado PDF con SHA-256.",
    "Producto en estado 'PENDIENTE' con ingredientes registrados.",
    "1. Selecciona producto pendiente. 2. Sistema analiza cada ingrediente. "
    "3. Calcula Eco-Score restando puntos por sustancias nocivas. "
    "4. Eco-Score ≥ 70: 'APROBADO'; < 70: 'RECHAZADO'. "
    "5. Genera certificado PDF con resultados y hash SHA-256. 6. Actualiza estado del producto.",
    "Producto publicado si aprobado. Registro de auditoría inmutable con hash de verificación.")

use_case(doc,
    "CU-05", "Compra de Productos",
    "Cliente",
    "Permite al Cliente agregar productos al carrito y completar la compra con métodos de pago locales "
    "(Yape, Plin, TuPay) o internacionales (Stripe).",
    "Cliente autenticado. Productos aprobados con stock disponible.",
    "1. Navega el catálogo y agrega al carrito. 2. Procede al checkout. "
    "3. Selecciona método de envío y pago. 4. Para locales: genera código de referencia y confirma. "
    "5. Para Stripe: redirige a pasarela segura. 6. Confirma pago, crea la orden y descuenta stock.",
    "Orden con estado 'CONFIRMADO'. Stock descontado. Notificaciones enviadas a productor y cliente.")

use_case(doc,
    "CU-06", "Ver Catálogo y Filtrar Productos",
    "Usuario anónimo, Cliente",
    "Permite a cualquier visitante navegar el catálogo de productos aprobados con filtros por "
    "categoría, texto y Eco-Score mínimo, y ver el detalle completo de cada producto.",
    "No se requiere autenticación. Solo se muestran productos con estado 'APROBADO'.",
    "1. Accede a la página principal. 2. Ve productos destacados. 3. Aplica filtros. "
    "4. Selecciona producto para ver detalle: descripción, ingredientes, origen GPS, certificado PDF, "
    "Eco-Score y badges.",
    "Usuario visualiza información completa del producto incluyendo resultados de auditoría química.")

use_case(doc,
    "CU-07", "Gestión de Usuarios",
    "Administrador",
    "Permite gestionar cuentas: activar/desactivar usuarios, verificar el RUC de Productores "
    "y sus certificaciones ecológicas, y consultar el historial de actividades.",
    "Administrador autenticado con rol ADMIN.",
    "1. Accede al panel de usuarios. 2. Filtra por rol y estado. 3. Selecciona usuario. "
    "4. Activa/desactiva cuenta, verifica RUC, otorga certificación o elimina la cuenta.",
    "Usuario actualizado con los cambios del Administrador. Log de auditoría registrado.")

use_case(doc,
    "CU-08", "Generación de Certificado PDF de Auditoría",
    "Administrador, Cliente",
    "Genera y descarga el certificado de auditoría química en PDF: Eco-Score, tabla de ingredientes "
    "con nivel de riesgo, badges obtenidos, estado final y hash SHA-256.",
    "El producto debe haber sido auditado. El usuario debe tener acceso al certificado.",
    "1. Solicita certificado del producto auditado. 2. Sistema consulta resultados almacenados. "
    "3. Genera PDF: logo EcoMarket, datos del producto, tabla de ingredientes, Eco-Score visual, "
    "badges y hash SHA-256 de verificación.",
    "PDF descargado con hash SHA-256 que certifica la autenticidad e inmutabilidad de la auditoría.")

use_case(doc,
    "CU-09", "Dashboard de Ventas y Estadísticas",
    "Productor, Administrador",
    "Panel con gráficos de ventas, productos más vendidos, ingresos por período y métricas de auditoría. "
    "El Productor ve solo sus datos; el Administrador ve datos globales del sistema.",
    "Usuario autenticado con rol PRODUCTOR o ADMIN.",
    "1. Accede al dashboard. 2. Visualiza gráficos: diario, semanal, mensual. "
    "3. Revisa productos más vendidos, ingresos totales y comisiones. "
    "4. Administrador: total de usuarios, productos auditados, tasa de aprobación/rechazo.",
    "Usuario obtiene visión clara del rendimiento de ventas y operaciones del sistema.")

use_case(doc,
    "CU-10", "Notificaciones Push",
    "Sistema (automático), Productor, Cliente",
    "El sistema envía notificaciones automáticas ante eventos: pago confirmado, "
    "cambio de estado de producto, resultado de auditoría y actualización de pedido.",
    "Evento relevante ocurrido en el sistema.",
    "1. Ocurre un evento (pago, auditoría, cambio de estado). 2. Sistema prepara notificación con "
    "título, cuerpo y datos adicionales. 3. Envía al usuario destino. 4. Usuario recibe notificación.",
    "Usuario notificado en tiempo real. Mejora la experiencia y reduce el tiempo de respuesta.")

doc.add_page_break()

# ── 4. GLOSARIO ──────────────────────────────────────────────
heading(doc, "4. Glosario de Términos", 1)

glosario = [
    ("Eco-Score",
     "Puntuación del 0 al 100 que indica el nivel de pureza ecológica de un producto. "
     "Se calcula restando puntos por cada ingrediente nocivo detectado en la auditoría. Puntaje ≥ 70 = aprobado."),
    ("Badge",
     "Insignia visual otorgada según resultados de auditoría: Eco-Friendly, Toxic-Free, Vegan, Cruelty Free y Plastic Free."),
    ("Auditoría Química",
     "Proceso automatizado mediante IA que analiza ingredientes contra una BD de sustancias prohibidas, "
     "evaluando su impacto en la salud y el medio ambiente."),
    ("JWT (JSON Web Token)",
     "Token de autenticación estándar que permite acceder a recursos protegidos sin reenviar credenciales en cada solicitud."),
    ("Yape / Plin",
     "Métodos de pago móvil populares en Perú integrados en EcoMarket para usuarios sin tarjetas de crédito."),
    ("Stripe",
     "Pasarela de pago internacional para compradores con tarjetas de crédito/débito, usada para ventas al exterior."),
    ("SHA-256",
     "Algoritmo criptográfico de 256 bits que genera un identificador único e inmutable para cada certificado de auditoría."),
    ("RUC",
     "Registro Único de Contribuyentes. Número de identificación fiscal en Perú requerido para Productores."),
    ("SENASA",
     "Servicio Nacional de Sanidad Agraria. Entidad peruana que certifica productos orgánicos y ecológicos."),
    ("Microservicios",
     "Arquitectura donde cada función (auth, pagos, auditoría, IA, notificaciones) es un servicio independiente "
     "con su propia base de datos, escalable y desplegable de forma autónoma."),
    ("API Gateway",
     "Punto de entrada único al backend que enruta las solicitudes del frontend a cada microservicio, "
     "gestionando CORS, seguridad y balanceo de carga."),
]
for term, defn in glosario:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"{term}: "); font(r1, bold=True, color=C_GREEN, size=10)
    r2 = p.add_run(defn);        font(r2, size=10, color=C_DARK)

# ── Guardar ──────────────────────────────────────────────────
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_WORD = os.path.join(ROOT, "docs", "EcoMarket_Documentacion_Profesional.docx")
doc.save(OUT_WORD)
print(f"[OK] Word guardado en: {OUT_WORD}")

# ════════════════════════════════════════════════════════════
# ZIP DEL PROYECTO
# ════════════════════════════════════════════════════════════
EXCLUDE_DIRS = {
    ".git", "node_modules", "__pycache__", ".next", "target",
    ".venv", "venv", "env", ".mvn", ".idea", ".vs", "dist",
    "build", ".gradle",
}
EXCLUDE_FILES = {
    ".DS_Store", "Thumbs.db", ".gitkeep",
}
EXCLUDE_EXT = {".pyc", ".class", ".log", ".tmp", ".cache"}
# Include .env.example but not .env
INCLUDE_DOTFILES = {".env.example", ".gitignore", ".dockerignore"}

OUT_ZIP = os.path.join(ROOT, "EcoMarket-Software.zip")
file_count = 0

with zipfile.ZipFile(OUT_ZIP, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as zf:
    for dirpath, dirnames, filenames in os.walk(ROOT):
        # Prune excluded dirs in-place
        dirnames[:] = [
            d for d in dirnames
            if d not in EXCLUDE_DIRS and not d.startswith(".")
        ]
        for filename in filenames:
            # Skip hidden files except allowed ones
            if filename.startswith(".") and filename not in INCLUDE_DOTFILES:
                continue
            if filename in EXCLUDE_FILES:
                continue
            ext = os.path.splitext(filename)[1].lower()
            if ext in EXCLUDE_EXT:
                continue
            filepath = os.path.join(dirpath, filename)
            arcname = os.path.relpath(filepath, ROOT)
            zf.write(filepath, arcname)
            file_count += 1

size_mb = os.path.getsize(OUT_ZIP) / (1024 * 1024)
print(f"[OK] ZIP guardado en: {OUT_ZIP}")
print(f"     {file_count} archivos · {size_mb:.1f} MB")
